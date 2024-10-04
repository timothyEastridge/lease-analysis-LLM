import streamlit as st
import os
import tempfile
import shutil
from typing import List, Dict
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.chat_models import ChatOpenAI
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain.text_splitter import RecursiveCharacterTextSplitter
from openai import OpenAI, APIError, RateLimitError, APIConnectionError, APITimeoutError
import traceback
import openai
import io
import zipfile
import tiktoken

# Set the page configuration
st.set_page_config(layout='wide', page_title="Lease Synopsis Generator", page_icon="📄")

# Custom CSS for consistent styling
st.markdown("""
<style>
    .stButton > button {
        width: 100%;
        background-color: #4CAF50;
        color: white !important;
        font-size: 18px;
        font-weight: bold;
        padding: 10px 24px;
        border-radius: 5px;
        border: none;
        cursor: pointer;
        transition: background-color 0.3s ease;
    }
    .stButton > button:hover {
        background-color: #45a049;
    }
    .stTextInput > div > div > input {
        font-size: 16px;
    }
    h1 {
        color: #2C3E50;
        text-align: center;
        padding-bottom: 20px;
    }
    h2 {
        color: #34495E;
    }
    .fullWidth {
        width: 100%;
    }
    .reportview-container .main .block-container {
        max-width: 95%;
        padding-top: 5rem;
        padding-right: 1rem;
        padding-left: 1rem;
        padding-bottom: 5rem;
    }
    .document-preview {
        border: 1px solid #ddd;
        padding: 10px;
        margin-bottom: 10px;
        border-radius: 5px;
    }
    .document-preview h3 {
        margin-top: 0;
        color: #2C3E50;
    }
</style>
""", unsafe_allow_html=True)

# App Title and Intro
st.title("📄 Lease Synopsis Generator and Chatbot")
st.markdown("---")

# OpenAI API key setup using Streamlit secrets
if "openai" not in st.secrets or "api_key" not in st.secrets["openai"]:
    st.error("OpenAI API key not found in Streamlit secrets. Please add it to your app's secrets under [openai] api_key.")
    st.stop()

openai_api_key = st.secrets["openai"]["api_key"]
os.environ["OPENAI_API_KEY"] = openai_api_key
openai.api_key = openai_api_key

# Initialize session state variables
if 'chat_history' not in st.session_state:
    st.session_state.chat_history: List[tuple] = []
if 'vector_store' not in st.session_state:
    st.session_state.vector_store = None
if 'document_previews' not in st.session_state:
    st.session_state.document_previews: Dict[str, str] = {}

# Functions from improved_lease_synopsis_generator.py
def num_tokens_from_string(string: str, encoding_name: str = "cl100k_base") -> int:
    encoding = tiktoken.get_encoding(encoding_name)
    num_tokens = len(encoding.encode(string))
    return num_tokens

def chunk_document(text, max_tokens=4000, chunk_overlap=200):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=max_tokens,
        chunk_overlap=chunk_overlap,
        length_function=lambda x: num_tokens_from_string(x)
    )
    return text_splitter.split_text(text)

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    return '\n'.join([para.text for para in doc.paragraphs])

def create_formatted_docx(content, output_file):
    doc = docx.Document()
    
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    
    header_style = doc.styles.add_style('Header Style', WD_STYLE_TYPE.PARAGRAPH)
    header_style.font.name = 'Calibri'
    header_style.font.size = Pt(16)
    header_style.font.bold = True
    header_style.font.color.rgb = RGBColor(0, 0, 128)  # Navy Blue
    
    header = doc.add_paragraph("Lease Synopsis", style='Header Style')
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    
    table.columns[0].width = Cm(7)
    table.columns[1].width = Cm(12)
    
    lines = content.split('\n')
    for line in lines:
        if ':' in line:
            key, value = line.split(':', 1)
            row_cells = table.add_row().cells
            row_cells[0].text = key.strip()
            row_cells[1].text = value.strip().replace('*', '')
            
            key_para = row_cells[0].paragraphs[0]
            key_para.runs[0].bold = True
            key_para.runs[0].font.color.rgb = RGBColor(0, 0, 128)
            
            value_para = row_cells[1].paragraphs[0]
            value_para.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    doc.save(output_file)

def generate_response(document_text, chat_llm):
    prompt_template = """
    You are an expert in lease document analysis. Given a lease document or a portion of a lease document, extract the following information. Be concise and only return the information requested of you. If information is not found or not applicable, write "" instead of "N/A". Do not include duplicate or redundant information.
    
    Please fill out the following template:
    
    Location:
    Address:
    Tenant Reference Name (doing business as):
    Tenant Entity:
    Guarantor:
    Tenant's Notice Address (prior to occupancy):
    Tenant's Notice Address (after occupancy):
    Landlord's Notice Address (if mailed):
    Landlord's Notice Address (if delivered):
    Landlord's Payment Address:
    Leased Premises:
    Square Feet:
    Commencement Date:
    Expiration Date:
    Extension Options:
    Base Rent:
    Operating Expenses:
    Parking:
    Construction/Allowance:
    Landlord's Relocation Rights:
    Tenant's Preferential Rights:
    Termination Options:
    Sign Rights:
    Exclusive:
    Use Restrictions on Landlord:
    Build Restrictions on Landlord:
    Off-site restrictions on Landlord:
    Security Deposit:
    Default Cure Period:
    Holdover:
    Broker/Commission:
    Notice Address:
    Other Provisions:
    Hazardous Material:
    Insurance:
    Tenant's Broker:
    Special Provisions:
    
    Document Text: {document_text}
    """

    prompt = PromptTemplate(template=prompt_template, input_variables=["document_text"])
    chain = LLMChain(llm=chat_llm, prompt=prompt)

    try:
        chunks = chunk_document(document_text)
        responses = []
        for chunk in chunks:
            response_dict = chain.invoke({"document_text": chunk})
            responses.append(response_dict.get('text', '').strip())
        
        combined_response = '\n'.join(responses)
        final_response = post_process_response(combined_response)
        
        return final_response
    except Exception as e:
        return f'Error in processing: {str(e)}'

def post_process_response(response):
    expected_fields = [
        "Location",
        "Address",
        "Tenant Reference Name (doing business as)",
        "Tenant Entity",
        "Guarantor",
        "Tenant's Notice Address (prior to occupancy)",
        "Tenant's Notice Address (after occupancy)",
        "Landlord's Notice Address (if mailed)",
        "Landlord's Notice Address (if delivered)",
        "Landlord's Payment Address",
        "Leased Premises",
        "Square Feet",
        "Commencement Date",
        "Expiration Date",
        "Extension Options",
        "Base Rent",
        "Operating Expenses",
        "Parking",
        "Construction/Allowance",
        "Landlord's Relocation Rights",
        "Tenant's Preferential Rights",
        "Termination Options",
        "Sign Rights",
        "Exclusive",
        "Use Restrictions on Landlord",
        "Build Restrictions on Landlord",
        "Off-site restrictions on Landlord",
        "Security Deposit",
        "Default Cure Period",
        "Holdover",
        "Broker/Commission",
        "Notice Address",
        "Other Provisions",
        "Hazardous Material",
        "Insurance",
        "Tenant's Broker",
        "Special Provisions"
    ]
    lines = response.split('\n')
    processed_fields = {field: "" for field in expected_fields}
    
    for line in lines:
        if ':' in line:
            key, value = line.split(':', 1)
            key = key.strip()
            value = value.strip().replace('*', '')
            if key in processed_fields:
                if processed_fields[key]:
                    processed_fields[key] += f"; {value}"
                else:
                    processed_fields[key] = value
            else:
                # Handle unexpected fields if necessary
                pass
    
    final_response = '\n'.join([f"{key}: {processed_fields[key]}" for key in expected_fields])
    return final_response


def process_docx_files(folder_path):
    docx_files = [file for file in os.listdir(folder_path) if file.endswith('.docx')]
    reports_folder = os.path.join(folder_path, 'Reports')
    os.makedirs(reports_folder, exist_ok=True)
    
    chat_llm = ChatOpenAI(temperature=0.1, model="gpt-4")
    
    all_synopses = []
    
    for docx_file in docx_files:
        docx_path = os.path.join(folder_path, docx_file)
        document_text = extract_text_from_docx(docx_path)
        
        if document_text.strip():
            response = generate_response(document_text, chat_llm)
            output_file = os.path.join(reports_folder, f"synopsis_{os.path.splitext(docx_file)[0]}.docx")
            create_formatted_docx(response, output_file)
            all_synopses.append(response)
    
    # Generate consolidated report
    consolidated_report = consolidate_synopses(all_synopses)
    summarized_report = summarize_consolidated_synopsis(consolidated_report, chat_llm)
    
    create_formatted_docx(summarized_report, os.path.join(reports_folder, "consolidated_synopsis.docx"))
    
    return reports_folder

def consolidate_synopses(all_synopses):
    consolidated = {}
    for synopsis in all_synopses:
        lines = synopsis.split('\n')
        for line in lines:
            if ':' in line:
                key, value = line.split(':', 1)
                key = key.strip()
                value = value.strip().replace('*', '')
                if key not in consolidated:
                    consolidated[key] = value
                elif value and value != consolidated[key] and value != "Not specified":
                    consolidated[key] += f"; {value}"
    
    return '\n'.join([f"{key}: {value}" for key, value in consolidated.items()])

def summarize_consolidated_synopsis(consolidated_synopsis, chat_llm):
    summary_prompt_template = """
    You are an expert in lease document analysis. Be concise and only return the information requested of you. Given the following concatenated lease document information, simplify the redundant parts of the text but maintain all the relevant detail:
    Consolidated Information: {consolidated_synopsis}
    """
    
    summary_prompt = PromptTemplate(template=summary_prompt_template, input_variables=["consolidated_synopsis"])
    summary_chain = LLMChain(llm=chat_llm, prompt=summary_prompt)
    
    chunks = chunk_document(consolidated_synopsis, max_tokens=8000)
    summarized_chunks = []
    
    for chunk in chunks:
        summarized_chunk = summary_chain.invoke({"consolidated_synopsis": chunk}).get('text', '').strip()
        summarized_chunks.append(summarized_chunk)
    
    return '\n\n'.join(summarized_chunks)

def process_uploaded_files(uploaded_files) -> None:
    """Process uploaded files and generate lease synopses."""
    with tempfile.TemporaryDirectory() as temp_dir:
        # Save uploaded files and extract text
        documents = []
        for uploaded_file in uploaded_files:
            file_path = os.path.join(temp_dir, uploaded_file.name)
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            text = extract_text_from_docx(file_path)
            documents.append(text)
            
            # Create document preview
            preview = text[:500] + "..." if len(text) > 500 else text
            st.session_state.document_previews[uploaded_file.name] = preview
        
        # Generate Lease Synopses
        reports_folder = process_docx_files(temp_dir)
        
        # Zip the synopses for download
        zip_path = shutil.make_archive(os.path.join(temp_dir, "lease_synopses"), 'zip', reports_folder)
        
        st.success("✅ Lease synopses generated successfully")
        
        # Download button for the generated synopses
        with open(zip_path, "rb") as file:
            st.download_button(
                label="Download Lease Synopses",
                data=file,
                file_name="lease_synopses.zip",
                mime="application/zip"
            )
        
        # Create the vector store using FAISS
        embeddings = OpenAIEmbeddings()
        st.session_state.vector_store = FAISS.from_texts(documents, embeddings)
        
        st.success("✅ Chatbot prepared successfully")

def handle_user_input(user_question: str, conversation_chain) -> None:
    """Process user input and generate AI response."""
    try:
        # Split the user question if it's too long
        question_chunks = chunk_document(user_question, max_tokens=4000)
        responses = []
        
        for chunk in question_chunks:
            response = conversation_chain({"question": chunk})
            responses.append(response['answer'])
        
        combined_response = ' '.join(responses)
        
        st.session_state.chat_history.append(("You", user_question))
        st.session_state.chat_history.append(("AI", combined_response))
    except (APIError, RateLimitError, APIConnectionError, APITimeoutError) as e:
        st.error(f"An error occurred with the OpenAI API: {str(e)}")
    except Exception as e:
        st.error(f"An unexpected error occurred: {str(e)}")
        st.error(traceback.format_exc())

def show_file_preview(uploaded_file):
    """Display a preview of the uploaded file."""
    if uploaded_file is not None:
        st.write("File Preview:")
        doc = docx.Document(uploaded_file)
        for para in doc.paragraphs[:5]:  # Show first 5 paragraphs
            st.write(para.text)
        st.write("...")

def main():
    # File Upload Section
    uploaded_files = st.file_uploader("Upload .docx files", type="docx", accept_multiple_files=True)

    if uploaded_files:
        st.success(f"✅ {len(uploaded_files)} file(s) uploaded successfully")
        

        
        if st.button("Generate Lease Synopsis and Prepare Chatbot"):
            with st.spinner("🔎 Generating lease synopsis and preparing chatbot..."):
                try:
                    process_uploaded_files(uploaded_files)
                except Exception as e:
                    st.error(f"An error occurred: {str(e)}")
                    st.error("Stack trace:", exc_info=True)
        
        # Chatbot Interface
        if st.session_state.vector_store is not None:
            st.subheader("Chat with your Lease Documents")
            
            # Initialize the Conversational Chain
            llm = ChatOpenAI(temperature=0)
            memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)
            conversation_chain = ConversationalRetrievalChain.from_llm(
                llm=llm,
                retriever=st.session_state.vector_store.as_retriever(),
                memory=memory
            )
            
            # User Input for Chatbot
            user_question = st.text_input("Ask a question about your lease documents:")
            if user_question:
                handle_user_input(user_question, conversation_chain)
            
            # Display the chat history
            for role, message in st.session_state.chat_history:
                if role == "You":
                    st.write(f"👤 **You:** {message}")
                else:
                    st.write(f"🤖 **AI:** {message}")

    else:
        st.warning("⚠️ Please upload .docx files to generate lease synopses and prepare the chatbot")

    st.markdown("<br>" * 15, unsafe_allow_html=True)
    st.markdown("---")
    st.markdown("Created with ❤️ by Weaver")

if __name__ == "__main__":
    main()
