# import os
# import keyring
# import openai
# import docx
# from docx.shared import Pt, Cm, RGBColor
# from docx.enum.style import WD_STYLE_TYPE
# from docx.enum.table import WD_TABLE_ALIGNMENT
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from langchain.chat_models import ChatOpenAI
# from langchain.prompts import PromptTemplate
# from langchain.chains import LLMChain
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from tqdm import tqdm
# from tenacity import retry, stop_after_attempt, wait_exponential, RetryError
# import concurrent.futures

# os.environ['OPENAI_API_KEY'] = keyring.get_password('eastridge', 'openai')
# openai.api_key = os.getenv('OPENAI_API_KEY')

# @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
# def create_chat_llm():
#     return ChatOpenAI(temperature=0.1, model="gpt-4o")

# chat_llm = create_chat_llm()

# prompt_template = """
# You are an expert in lease document analysis. Given a lease document or a portion of a lease document, extract the following information. Be concise and specific. If information is not found or not applicable, write "Not specified" instead of "N/A". Do not include duplicate or redundant information.

# Tenant Reference Name (doing business as):
# Tenant Entity:
# Guarantor: (keep this concise. For example, instead of Cenovus Energy Inc., a Canadian corporation just say Cenovus Energy Inc.)
# List of Documents:
# Leased Premises:
# Square Feet:
# Commencement Date:
# Expiration Date:
# Extension Options:
# Base Rent:
# Operating Expenses:
# Parking:
# Construction/Allowance:
# Landlord's Relocation Rights:
# Tenant's Preferential Rights:
# Termination Options:
# Sign Rights:
# Exclusive:
# Use Restrictions on Landlord:
# Build Restrictions on Landlord:
# Off-site restrictions on Landlord:
# Security Deposit:
# Default Cure Period:
# Holdover:
# Broker/Commission:
# Notice Address:
# Other Provisions:

# Document Text: {document_text}
# """

# prompt = PromptTemplate(template=prompt_template, input_variables=["document_text"])
# chain = LLMChain(llm=chat_llm, prompt=prompt)

# def chunk_document(text, chunk_size=6000, chunk_overlap=200):
#     text_splitter = RecursiveCharacterTextSplitter(
#         chunk_size=chunk_size,
#         chunk_overlap=chunk_overlap,
#         length_function=len
#     )
#     return text_splitter.split_text(text)

# @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
# def process_chunk(chunk):
#     try:
#         response_dict = chain.invoke({"document_text": chunk})
#         return response_dict.get('text', '').strip()
#     except Exception as e:
#         return f'Error in processing chunk: {str(e)}'

# def generate_response(document_text):
#     try:
#         chunks = chunk_document(document_text)
        
#         with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
#             responses = list(executor.map(process_chunk, chunks))
        
#         combined_response = '\n'.join(responses)
#         final_response = post_process_response(combined_response)
        
#         return final_response
#     except Exception as e:
#         return f'Error in processing: {str(e)}'

# def post_process_response(response):
#     lines = response.split('\n')
#     processed_fields = {}
    
#     for line in lines:
#         if ':' in line:
#             key, value = line.split(':', 1)
#             key = key.strip()
#             value = value.strip()
            
#             if key not in processed_fields:
#                 processed_fields[key] = value
#             elif value and value != processed_fields[key] and value != "Not specified":
#                 processed_fields[key] += f"; {value}"
    
#     final_response = '\n'.join([f"{key}: {value}" for key, value in processed_fields.items()])
    
#     return final_response

# def extract_text_from_docx(docx_path):
#     doc = docx.Document(docx_path)
#     return '\n'.join([para.text for para in doc.paragraphs])

# def process_docx_files(folder_path):
#     docx_files = [file for file in os.listdir(folder_path) if file.endswith('.docx')]
#     reports_folder = os.path.join(folder_path, 'Reports')
#     os.makedirs(reports_folder, exist_ok=True)
    
#     all_synopses = []
    
#     with tqdm(total=len(docx_files), desc="Processing lease documents") as pbar:
#         for docx_file in docx_files:
#             docx_path = os.path.join(folder_path, docx_file)
#             document_text = extract_text_from_docx(docx_path)
            
#             if document_text.strip():
#                 response = generate_response(document_text)
#                 output_file = os.path.join(reports_folder, f"synopsis_{os.path.splitext(docx_file)[0]}.docx")
#                 create_formatted_docx(response, output_file)
#                 all_synopses.append(response)
#             pbar.update(1)
    
#     # Generate consolidated report
#     consolidated_report = consolidate_synopses(all_synopses)
#     create_formatted_docx(consolidated_report, os.path.join(reports_folder, "consolidated_synopsis.docx"))
    
#     return reports_folder

# def create_formatted_docx(content, output_file):
#     doc = docx.Document()
    
#     normal_style = doc.styles['Normal']
#     normal_style.font.name = 'Calibri'
#     normal_style.font.size = Pt(11)
    
#     header_style = doc.styles.add_style('Header Style', WD_STYLE_TYPE.PARAGRAPH)
#     header_style.font.name = 'Calibri'
#     header_style.font.size = Pt(16)
#     header_style.font.bold = True
#     header_style.font.color.rgb = RGBColor(0, 0, 128)  # Navy Blue
    
#     header = doc.add_paragraph("Lease Synopsis", style='Header Style')
#     header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
#     table = doc.add_table(rows=1, cols=2)
#     table.style = 'Table Grid'
#     table.autofit = False
#     table.allow_autofit = False
    
#     table.columns[0].width = Cm(7)
#     table.columns[1].width = Cm(12)
    
#     lines = content.split('\n')
#     for line in lines:
#         if ':' in line:
#             key, value = line.split(':', 1)
#             row_cells = table.add_row().cells
#             row_cells[0].text = key.strip()
#             row_cells[1].text = value.strip()
            
#             key_para = row_cells[0].paragraphs[0]
#             key_para.runs[0].bold = True
#             key_para.runs[0].font.color.rgb = RGBColor(0, 0, 128)  # Navy Blue
            
#             value_para = row_cells[1].paragraphs[0]
#             value_para.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Black
    
#     table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
#     doc.save(output_file)

# def consolidate_synopses(all_synopses):
#     consolidated = {}
#     for synopsis in all_synopses:
#         lines = synopsis.split('\n')
#         for line in lines:
#             if ':' in line:
#                 key, value = line.split(':', 1)
#                 key = key.strip()
#                 value = value.strip()
#                 if key not in consolidated:
#                     consolidated[key] = value
#                 elif value and value != consolidated[key] and value != "Not specified":
#                     consolidated[key] += f"; {value}"
    
#     return '\n'.join([f"{key}: {value}" for key, value in consolidated.items()])

# if __name__ == "__main__":
#     folder_path = r'C:\\Users\\TimEa\\OneDrive\\Data\\Weaver'
#     reports_folder = process_docx_files(folder_path)
#     print(f"Completed processing. Individual synopses and consolidated report generated in the Reports folder.")

import os
import keyring
import openai
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from langchain.chat_models import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain.text_splitter import RecursiveCharacterTextSplitter
from tqdm import tqdm
from tenacity import retry, stop_after_attempt, wait_exponential, RetryError
import concurrent.futures

os.environ['OPENAI_API_KEY'] = keyring.get_password('eastridge', 'openai')
openai.api_key = os.getenv('OPENAI_API_KEY')

@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
def create_chat_llm():
    return ChatOpenAI(temperature=0.1, model="gpt-4o")

chat_llm = create_chat_llm()

prompt_template = """
You are an expert in lease document analysis. Given a lease document or a portion of a lease document, extract the following information. Be concise and specific. If information is not found or not applicable, write "" instead of "N/A". Do not include duplicate or redundant information.

Tenant Reference Name (doing business as):
Tenant Entity:
Guarantor: (keep this concise. For example, instead of Cenovus Energy Inc., a Canadian corporation just say Cenovus Energy Inc.)
List of Documents:
Leased Premises:
Square Feet:
Commencement Date:
Expiration Date:
Extension Options:
Base Rent:
Operating Expenses: (add as much information as possible. for example: Tenant pays proportionate share of expenses (net lease), grossed up to reflect 100% occupancy. Management fee is 4% of rents. Commencing 04/01/25, expenses are capped at 6% accumulating and compounding amounts over FYE 03/31/25 amounts, except as attributable to insurance premiums/deductibles, increases in security due to staffing levels, janitorial or other costs which increase due to unionization, utilities and real estate tax/protest costs.)
Parking: (add as much info as possible. For example, Up to 24 non-reserved @ $50 per month, of which up to 5 may be reserved @ $100 per month. At any point during the term, T may convert an additional 2 non-reserved parking spaces into reserved parking spaces**Non-reserved parking charges abate from 10/01/24 - 02/16/30.)
Construction/Allowance:
Landlord's Relocation Rights: (add as much info as possible. for example, Landlord may relocate Tenant once during the term (except during first 2 years and last year of initial term and except during the  first or last year of extension option) upon 120 days prior written to another space in the building on 14th floor or higher, of a size between 100% to 110% of the premises at LL cost. Substitute premises to be improved with reasonably comparable or better quality leasehold improvements as existed in the premises. LL will provide T with at least 30 days access to the substitute premises after LL's tender  in order for T to install wiring, cabling, furniture, fixtures and equipment in the substitute premises at no cost to Tenant (Sec.3.3))
Tenant's Preferential Rights:
Termination Options:
Sign Rights:
Exclusive:
Use Restrictions on Landlord:
Build Restrictions on Landlord:
Off-site restrictions on Landlord:
Security Deposit:
Default Cure Period:
Holdover:
Broker/Commission:
Notice Address: (be sure to look for corporate address. for example, Leased PremisesWith a copy of all notices of default to the Guarantor at: Cenovus Energy Inc. 225 6 Avenue SW Calgary, AB T 2P 1N2 Attn: Director, Enterprise Compliance & Credit email: creditgroup@cenovus.com copy to: downstream.legal@cenovus.com)
Other Provisions:

Document Text: {document_text}
"""

prompt = PromptTemplate(template=prompt_template, input_variables=["document_text"])
chain = LLMChain(llm=chat_llm, prompt=prompt)

def chunk_document(text, chunk_size=6000, chunk_overlap=200):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len
    )
    return text_splitter.split_text(text)

@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
def process_chunk(chunk):
    try:
        response_dict = chain.invoke({"document_text": chunk})
        return response_dict.get('text', '').strip()
    except Exception as e:
        return f'Error in processing chunk: {str(e)}'

def generate_response(document_text):
    try:
        chunks = chunk_document(document_text)
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
            responses = list(executor.map(process_chunk, chunks))
        
        combined_response = '\n'.join(responses)
        final_response = post_process_response(combined_response)
        
        return final_response
    except Exception as e:
        return f'Error in processing: {str(e)}'

def post_process_response(response):
    lines = response.split('\n')
    processed_fields = {}
    
    for line in lines:
        if ':' in line:
            key, value = line.split(':', 1)
            key = key.strip()
            value = value.strip()
            
            if key not in processed_fields:
                processed_fields[key] = value
            elif value and value != processed_fields[key] and value != "Not specified":
                processed_fields[key] += f"; {value}"
    
    final_response = '\n'.join([f"{key}: {value}" for key, value in processed_fields.items()])
    
    return final_response

def post_process_response(response):
    lines = response.split('\n')
    processed_fields = {}
    
    for line in lines:
        if ':' in line:
            key, value = line.split(':', 1)
            key = key.strip()
            value = value.strip().replace('*', '')  # Remove asterisks
            
            if key not in processed_fields:
                processed_fields[key] = value
            elif value and value != processed_fields[key] and value != "Not specified":
                processed_fields[key] += f"; {value}"
    
    final_response = '\n'.join([f"{key}: {value}" for key, value in processed_fields.items()])
    
    return final_response

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    return '\n'.join([para.text for para in doc.paragraphs])

def process_docx_files(folder_path):
    docx_files = [file for file in os.listdir(folder_path) if file.endswith('.docx')]
    reports_folder = os.path.join(folder_path, 'Reports')
    os.makedirs(reports_folder, exist_ok=True)
    
    all_synopses = []
    
    with tqdm(total=len(docx_files), desc="Processing lease documents") as pbar:
        for docx_file in docx_files:
            docx_path = os.path.join(folder_path, docx_file)
            document_text = extract_text_from_docx(docx_path)
            
            if document_text.strip():
                response = generate_response(document_text)
                output_file = os.path.join(reports_folder, f"synopsis_{os.path.splitext(docx_file)[0]}.docx")
                create_formatted_docx(response, output_file)
                all_synopses.append(response)
            pbar.update(1)
    
    # Generate consolidated report
    consolidated_report = consolidate_synopses(all_synopses)
    
    # NEW: Summarize the consolidated report
    summarized_report = summarize_consolidated_synopsis(consolidated_report)
    
    create_formatted_docx(summarized_report, os.path.join(reports_folder, "consolidated_synopsis.docx"))
    
    return reports_folder

def create_formatted_docx(content, output_file):
    doc = docx.Document()
    
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    
    header_style = doc.styles.add_style('Header Style', WD_STYLE_TYPE.PARAGRAPH)
    header_style.font.name = 'Calibri'
    header_style.font.size = Pt(16)
    header_style.font.bold = True
    header_style.font.color.rgb = RGBColor(0, 0, 128)  # Navy Blue
    
    header = doc.add_paragraph("Lease Synopsis", style='Header Style')
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    
    table.columns[0].width = Cm(7)
    table.columns[1].width = Cm(12)
    
    lines = content.split('\n')
    for line in lines:
        if ':' in line:
            key, value = line.split(':', 1)
            row_cells = table.add_row().cells
            row_cells[0].text = key.strip()
            row_cells[1].text = value.strip().replace('*', '')  # Remove asterisks
            
            key_para = row_cells[0].paragraphs[0]
            key_para.runs[0].bold = True
            key_para.runs[0].font.color.rgb = RGBColor(0, 0, 128)  # Navy Blue
            
            value_para = row_cells[1].paragraphs[0]
            value_para.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Black
    
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    doc.save(output_file)

def consolidate_synopses(all_synopses):
    consolidated = {}
    for synopsis in all_synopses:
        lines = synopsis.split('\n')
        for line in lines:
            if ':' in line:
                key, value = line.split(':', 1)
                key = key.strip()
                value = value.strip().replace('*', '')  # Remove asterisks
                if key not in consolidated:
                    consolidated[key] = value
                elif value and value != consolidated[key] and value != "Not specified":
                    consolidated[key] += f"; {value}"
    
    return '\n'.join([f"{key}: {value}" for key, value in consolidated.items()])

# NEW: Added function to summarize consolidated synopsis
def summarize_consolidated_synopsis(consolidated_synopsis):
    summary_prompt_template = """
    You are an expert in lease document analysis. Given the following concatenated lease document information, simplify the redundant parts of the text but maintain all the relevant detail:
    Consolidated Information: {consolidated_synopsis}
    """
    
    summary_prompt = PromptTemplate(template=summary_prompt_template, input_variables=["consolidated_synopsis"])
    summary_chain = LLMChain(llm=chat_llm, prompt=summary_prompt)
    
    return summary_chain.invoke({"consolidated_synopsis": consolidated_synopsis}).get('text', '').strip()

if __name__ == "__main__":
    folder_path = r'C:\\Users\\TimEa\\OneDrive\\Data\\Weaver'
    reports_folder = process_docx_files(folder_path)
    print(f"Completed processing. Individual synopses and consolidated report generated in the Reports folder.")